import requests
from bs4 import BeautifulSoup
from docx import Document
from urllib.parse import urljoin, urlparse
import os
import re

def sanitize_filename(filename):
    filename = filename.strip()
    filename = filename.replace(".docx", "")
    filename = re.sub(r'[^A-Za-z0-9_-]', '', filename)
    if not filename:
        filename = "output"
    return filename

def main(a, b):
    start_url = a
    file_name = sanitize_filename(b)

    output_dir = os.path.join(os.getcwd(), "crawled_links_docs")
    os.makedirs(output_dir, exist_ok=True)

    output_path = os.path.join(output_dir, f"{file_name}.docx")

    doc = Document()
    doc.add_heading('Crawled Data Report', 0)

    try:
        response = requests.get(start_url, timeout=10)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        print("Error fetching the page:", e)
        raise

    soup = BeautifulSoup(response.text, 'html.parser')

    links = []
    for tag in soup.find_all('a', href=True):
        full_url = urljoin(start_url, tag['href'])
        if urlparse(full_url).scheme in ['http', 'https']:
            links.append(full_url)

    doc.add_paragraph(f"Crawled URL: {start_url}")
    doc.add_paragraph(f"Number of links found: {len(links)}")

    doc.add_heading('Links:', level=1)
    for link in links:
        doc.add_paragraph(link, style='ListBullet')

    try:
        doc.save(output_path)
        print(f"✅ Crawled data saved to '{output_path}'")
    except Exception as e:
        print(f"❌ Failed to save Word document: {e}")
        raise
