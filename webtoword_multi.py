import requests
from bs4 import BeautifulSoup
from docx import Document
from urllib.parse import urljoin, urlparse
import os
import time
import re

def sanitize_filename(filename):
    filename = filename.strip()
    filename = filename.replace(".docx", "")
    filename = re.sub(r'[^A-Za-z0-9_-]', '', filename)
    if not filename:
        filename = "output"
    return filename

def main(a, b):
    max_pages_to_crawl = 10
    start_url = a
    file_name = sanitize_filename(b)

    output_dir = os.path.join(os.getcwd(), "crawled_links_docs")
    os.makedirs(output_dir, exist_ok=True)

    output_path = os.path.join(output_dir, f"{file_name}.docx")

    doc = Document()
    doc.add_heading('Crawled Data Report', 0)
    doc.add_paragraph(f'Start URL: {start_url}')
    doc.add_paragraph(f'Crawl started at: {time.strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('')

    to_crawl = [start_url]
    crawled = set()

    while to_crawl and len(crawled) < max_pages_to_crawl:
        url = to_crawl.pop(0)
        if url in crawled:
            continue

        print(f"Crawling: {url}")

        try:
            response = requests.get(url, timeout=10)  # Use url here (was start_url mistakenly)
            response.raise_for_status()
        except requests.exceptions.RequestException as e:
            print(f"Error fetching the page {url}: {e}")
            continue

        soup = BeautifulSoup(response.text, 'html.parser')
        title = soup.title.string.strip() if soup.title else 'No title'

        links = []
        for tag in soup.find_all('a', href=True):
            full_url = urljoin(url, tag['href'])
            parsed_url = urlparse(full_url)
            if parsed_url.scheme in ['http', 'https']:
                links.append(full_url)

        doc.add_heading(title, level=1)
        doc.add_paragraph(f'URL: {url}')
        doc.add_paragraph(f'Links found: {len(links)}')

        if links:
            doc.add_heading('Links:', level=2)
            for link in links:
                doc.add_paragraph(link, style='ListBullet')
        else:
            doc.add_paragraph('No links found.')

        doc.add_paragraph('')

        crawled.add(url)

        for link in links:
            if link not in crawled and link not in to_crawl:
                to_crawl.append(link)

    try:
        doc.save(output_path)
        print(f"✅ Crawled data saved to '{output_path}'")
    except Exception as e:
        print(f"❌ Failed to save Word document: {e}")
        raise
